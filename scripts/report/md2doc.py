#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2doc.py — 把 check_text.py 出嘅 .md 報告轉做 Word / Excel / HTML（已經 run 完就唔使重跑）。

點用（乜 flag 都唔使）：
    python scripts\\report\\md2doc.py                    # 轉 file_check\\_檢查報告\\*.md → .docx
    python scripts\\report\\md2doc.py --to xlsx          # 出一個 Excel（逐個檔一張 sheet，好過逐條 tick）
    python scripts\\report\\md2doc.py --to doc           # 出 .doc（HTML 版，Word 開得，唔使 python-docx）
    python scripts\\report\\md2doc.py --dir 某資料夾 --out 另一個資料夾
    python scripts\\report\\md2doc.py 單一個檔.md

支援嘅 markdown（check_text 出嗰種就夠）：# 標題、- 項目符號、| 表格 |、`code`、**粗體**。
唔係要做通用 markdown engine —— 淨係要俾項目組打得開、改得到。
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from tqdm import tqdm
except ImportError:            # 冇裝 tqdm 照跑
    def tqdm(it=None, **kw):
        return it
    tqdm.write = print

DEFAULT_DIR = "file_check/_檢查報告"
_INLINE = re.compile(r"`([^`]*)`|\*\*([^*]*)\*\*")
# XML 1.0 唔食嘅控制字元（pptx 入面嘅 \x0b 軟斷行、\x07 表格尾等）——
#   唔清就 lxml 會炒「All strings must be XML compatible」（2026-08-20 Melco/MGM 撞到）。
_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ufffe\uffff]")


def xml_safe(t):
    """控制字元 → 空格／刪走，令 python-docx / lxml 收得。"""
    return _CTRL.sub(lambda m: " " if m.group(0) in "\x0b\x0c" else "", str(t))


def _plain(t):
    """剝走 inline markdown（`code`、**bold**）→ 純文字，順手清控制字元。"""
    return xml_safe(_INLINE.sub(
        lambda m: m.group(1) if m.group(1) is not None else m.group(2), t))


def parse_md(text):
    """→ [(kind, payload)]；kind = h1/h2/h3 | li | p | table(list[list[str]])。"""
    blocks, rows = [], []
    # ⚠ 一定要【split 之前】清 —— str.splitlines() 會當 \x0b/\x0c 係換行，
    #   一行表格會俾佢劈開，之後認唔返做 | … | 就會靜靜咁唔見咗成行。
    text = xml_safe(text)

    def flush():
        nonlocal rows
        if rows:
            blocks.append(("table", rows)); rows = []
    for ln in text.splitlines():
        s = ln.rstrip()
        if s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):     # |---|---| 分隔行
                continue
            rows.append([_plain(c) for c in cells]); continue
        flush()
        if not s.strip():
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", s)
        if m:
            blocks.append((f"h{len(m.group(1))}", _plain(m.group(2)))); continue
        if s.lstrip().startswith(("- ", "* ")):
            blocks.append(("li", _plain(s.lstrip()[2:]))); continue
        blocks.append(("p", _plain(s)))
    flush()
    return blocks


# ── .docx（真 Word 檔，有真表格）────────────────────────────────────────
def to_docx(blocks, out, title):
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    d = docx.Document()
    st = d.styles["Normal"]
    st.font.name = "Microsoft JhengHei"; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "Microsoft JhengHei")
    NAVY = RGBColor(0x00, 0x33, 0x8D)
    for kind, payload in blocks:
        if kind.startswith("h"):
            p = d.add_paragraph(); r = p.add_run(payload)
            r.bold = True; r.font.size = Pt({"h1": 16, "h2": 13, "h3": 11}[kind])
            r.font.color.rgb = NAVY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif kind == "li":
            d.add_paragraph(payload, style="List Bullet")
        elif kind == "p":
            d.add_paragraph(payload)
        else:
            t = d.add_table(rows=0, cols=len(payload[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(payload):
                cells = t.add_row().cells
                for ci, v in enumerate(row[:len(cells)]):
                    cells[ci].text = v
                    if ri == 0:
                        for pr in cells[ci].paragraphs:
                            for rr in pr.runs:
                                rr.bold = True; rr.font.color.rgb = NAVY
            d.add_paragraph()
    d.save(out)


# ── .doc（HTML；Word 直接開得，零 dependency）─────────────────────────────
def to_html(blocks, out, title):
    esc = (lambda s: s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    L = ["<html><head><meta charset='utf-8'><title>" + esc(title) + "</title><style>",
         "body{font-family:'Microsoft JhengHei',sans-serif;font-size:10.5pt}",
         "h1{color:#00338D;font-size:17pt}h2{color:#00338D;font-size:13pt}h3{font-size:11pt}",
         "table{border-collapse:collapse;width:100%;margin:6pt 0}",
         "td,th{border:1px solid #999;padding:3pt 5pt;font-size:9.5pt;vertical-align:top}",
         "th{background:#00338D;color:#fff}", "</style></head><body>"]
    ul = False
    for kind, payload in blocks:
        if kind != "li" and ul:
            L.append("</ul>"); ul = False
        if kind.startswith("h"):
            L.append(f"<{kind}>{esc(payload)}</{kind}>")
        elif kind == "li":
            if not ul:
                L.append("<ul>"); ul = True
            L.append(f"<li>{esc(payload)}</li>")
        elif kind == "p":
            L.append(f"<p>{esc(payload)}</p>")
        else:
            L.append("<table>")
            for ri, row in enumerate(payload):
                tag = "th" if ri == 0 else "td"
                L.append("<tr>" + "".join(f"<{tag}>{esc(v)}</{tag}>" for v in row) + "</tr>")
            L.append("</table>")
    if ul:
        L.append("</ul>")
    L.append("</body></html>")
    out.write_text("\n".join(L), encoding="utf-8-sig")


# ── .xlsx（逐個 md 一張 sheet；表格合埋一張，加返「第幾頁」欄）──────────────
def to_xlsx(items, out):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = openpyxl.Workbook(); wb.remove(wb.active)
    hdr_f = Font(bold=True, color="FFFFFF")
    hdr_b = PatternFill("solid", fgColor="00338D")
    for name, blocks in items:
        ws = wb.create_sheet(re.sub(r"[\\/*?:\[\]]", "_", name)[:31])
        page, wrote_hdr, r = "", False, 1
        for kind, payload in blocks:
            if kind == "h2":
                page = payload; continue
            if kind != "table":
                continue
            for ri, row in enumerate(payload):
                if ri == 0:
                    if wrote_hdr:
                        continue
                    ws.append(["頁"] + row); wrote_hdr = True
                    for c in ws[1]:
                        c.font = hdr_f; c.fill = hdr_b
                    r = 1; continue
                ws.append([page] + row); r += 1
        for col, w in zip("ABCDEFG", (14, 14, 10, 34, 26, 34, 8)):
            ws.column_dimensions[col].width = w
        for row in ws.iter_rows(min_row=2):
            for c in row:
                c.alignment = Alignment(vertical="top", wrap_text=True)
        ws.freeze_panes = "A2"
        if not wrote_hdr:
            ws["A1"] = "✅ 冇揾到問題"
    wb.save(out)


def main():
    ap = argparse.ArgumentParser(description="把 check_text 出嘅 md 報告轉做 Word / Excel / HTML")
    ap.add_argument("src", nargs="?", default=DEFAULT_DIR, help=f"md 檔或資料夾（預設 {DEFAULT_DIR}）")
    ap.add_argument("--to", default="docx", choices=["docx", "doc", "html", "xlsx"])
    ap.add_argument("--out", default=None, help="出邊度（預設同 md 一齊）")
    a = ap.parse_args()

    src = Path(a.src)
    mds = [src] if src.is_file() else sorted(src.rglob("*.md"))     # 資料夾＝連子資料夾一齊掃
    if not mds:
        print(f"✗ {src.resolve()} 入面冇 .md（先跑 check_text.py）"); return
    outdir = Path(a.out) if a.out else (src if src.is_dir() else src.parent)
    outdir.mkdir(parents=True, exist_ok=True)
    print(f"── {src.resolve()}：{len(mds)} 個 md → {a.to}")

    parsed, bad = [], []
    for p in tqdm(mds, desc="讀 md", unit="檔", ncols=76):
        try:
            parsed.append((p.stem, parse_md(p.read_text(encoding="utf-8", errors="replace"))))
        except Exception as e:
            bad.append((p.name, str(e)[:120]))
    if a.to == "xlsx":
        o = outdir / "用字檢查.xlsx"
        to_xlsx(parsed, o)
        print(f"✓ {o.resolve()}（{len(parsed)} 張 sheet）")
    else:
        ok = 0
        for name, blocks in tqdm(parsed, desc=f"寫 {a.to}", unit="檔", ncols=76):
            o = outdir / f"{name}.{'doc' if a.to == 'doc' else a.to}"
            try:
                (to_docx if a.to == "docx" else to_html)(blocks, o, name)
            except ImportError:
                o = outdir / f"{name}.doc"; to_html(blocks, o, name)
                tqdm.write("  ⚠ 冇 python-docx → 出咗 .doc（HTML 版，Word 一樣開得）")
            except Exception as e:      # ⚠ 一個檔炒唔可以拖冧成批
                bad.append((name, str(e)[:120])); continue
            ok += 1
        print(f"✓ {ok}/{len(parsed)} 個檔寫咗落 {outdir.resolve()}")
    for n, e in bad:
        print(f"  ✗ {n}：{e}")


if __name__ == "__main__":
    main()
